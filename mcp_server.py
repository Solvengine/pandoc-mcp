#!/usr/bin/env python3
"""
MCP-Server für Pandoc (Dokumentkonvertierung)
Startet als STDIO-Server für Claude Code / Desktop.

Benötigt:
  pip install fastmcp pypandoc python-dotenv
  pandoc muss installiert sein (brew install pandoc oder apt install pandoc)
"""

import os
import logging
from pathlib import Path
from dotenv import load_dotenv
from mcp.server.fastmcp import FastMCP
import pypandoc

# -----------------------------------------------------------------------------
# Grund-Setup
# -----------------------------------------------------------------------------
logging.basicConfig(level=logging.INFO)
logging.info("Launching Pandoc MCP…")

load_dotenv()
mcp = FastMCP("Pandoc")

# -----------------------------------------------------------------------------
# Tools
# -----------------------------------------------------------------------------
@mcp.tool()
def convert_html_to_docx(
    input_html: str,
    output_path: str,
    header_text: str = None,
    footer_text: str = None,
    reference_doc: str = None,
    enable_toc: bool = False,
    toc_depth: int = 3,
    process_page_breaks: bool = True
) -> str:
    """
    Konvertiert HTML zu Word (DOCX) mit optionalen Header/Footer, TOC und Page Breaks.

    Args:
        input_html: HTML-String oder Pfad zur HTML-Datei
        output_path: Pfad für die Ausgabe-DOCX-Datei
        header_text: Optionaler Header-Text
        footer_text: Optionaler Footer-Text
        reference_doc: Pfad zu einer Referenz-DOCX-Datei für Styling (Standard: templates/solvengine.docx)
        enable_toc: Aktiviert automatisches Inhaltsverzeichnis
        toc_depth: Tiefe des Inhaltsverzeichnisses (1-6, Standard: 3)
        process_page_breaks: Konvertiert HTML page-break Marker zu Word Page Breaks

    Returns:
        Bestätigung mit Dateipfad
    """
    extra_args = []

    # Header und Footer als Metadaten hinzufügen
    if header_text:
        extra_args.extend(['--metadata', f'header-includes={header_text}'])

    if footer_text:
        extra_args.extend(['--metadata', f'footer-includes={footer_text}'])

    # Referenz-Dokument für Styling
    if reference_doc is None:
        # Standard-Template verwenden
        default_template = os.path.join(os.path.dirname(__file__), 'templates', 'solvengine.docx')
        if os.path.exists(default_template):
            reference_doc = default_template

    if reference_doc and os.path.exists(reference_doc):
        extra_args.extend(['--reference-doc', reference_doc])

    # Table of Contents
    if enable_toc:
        extra_args.append('--toc')
        extra_args.extend(['--toc-depth', str(toc_depth)])
        extra_args.append('--standalone')  # Required for TOC generation in DOCX

    # Page Break Verarbeitung
    if process_page_breaks:
        # Erlaubt raw HTML für Page Breaks
        extra_args.extend(['--from', 'html+raw_html'])

    # Input vorbereiten
    html_content = input_html
    if os.path.isfile(input_html):
        with open(input_html, 'r', encoding='utf-8') as f:
            html_content = f.read()

    # Page Breaks in HTML zu Pandoc-Markdown konvertieren
    if process_page_breaks:
        import re
        # Verschiedene Page Break Marker unterstützen
        html_content = re.sub(
            r'<div[^>]*class=["\']page-?break["\'][^>]*>.*?</div>',
            '\n\n\\newpage\n\n',
            html_content,
            flags=re.IGNORECASE | re.DOTALL
        )
        html_content = re.sub(
            r'<hr[^>]*class=["\']page-?break["\'][^>]*>',
            '\n\n\\newpage\n\n',
            html_content,
            flags=re.IGNORECASE
        )
        html_content = re.sub(
            r'<div[^>]*style=["\'][^"\']*page-break-after\s*:\s*always[^"\']*["\'][^>]*>.*?</div>',
            '\n\n\\newpage\n\n',
            html_content,
            flags=re.IGNORECASE | re.DOTALL
        )

    # Konvertierung durchführen
    output = pypandoc.convert_text(
        html_content,
        'docx',
        format='html',
        outputfile=output_path,
        extra_args=extra_args
    )

    logging.info(f"Converted HTML to DOCX: {output_path}")
    return f"Erfolgreich konvertiert nach: {output_path}"


@mcp.tool()
def convert_html_to_docx_with_template(
    input_html: str,
    output_path: str,
    template_path: str = None,
    enable_toc: bool = True,
    toc_depth: int = 3,
    process_page_breaks: bool = True,
    lua_filter_path: str = None,
    prevent_table_split: bool = True,
    toc_page_break: bool = True,
    doc_title: str = None,
    doc_subtitle: str = None,
    doc_author: str = None,
    doc_date: str = None,
    title_logo_path: str = None,
    title_logo_width_inches: float = 2.0
) -> str:
    """
    Konvertiert HTML zu Word (DOCX) mit einem Template, TOC und Page Breaks.
    Das Template definiert Header, Footer und Styling.

    Args:
        input_html: HTML-String oder Pfad zur HTML-Datei
        output_path: Pfad für die Ausgabe-DOCX-Datei
        template_path: Pfad zur Template-DOCX-Datei (Standard: templates/solvengine.docx)
        enable_toc: Aktiviert automatisches Inhaltsverzeichnis
        toc_depth: Tiefe des Inhaltsverzeichnisses (1-6, Standard: 3)
        process_page_breaks: Konvertiert HTML page-break Marker zu Word Page Breaks
        lua_filter_path: Pfad zu einem Lua-Filter (Standard: filters/table-style.lua für Tabellen-Styling)
        prevent_table_split: Verhindert dass Tabellen über Seitengrenzen getrennt werden
        toc_page_break: Fügt einen Seitenumbruch nach dem Inhaltsverzeichnis ein
        doc_title: Dokumenttitel für Titelseite
        doc_subtitle: Untertitel für Titelseite
        doc_author: Autor(en) für Titelseite
        doc_date: Datum für Titelseite (wenn None, wird aktuelles Datum verwendet)
        title_logo_path: Pfad zum Logo-Bild für Titelseite
        title_logo_width_inches: Breite des Logos in Zoll (Standard: 2.0)

    Returns:
        Bestätigung mit Dateipfad
    """
    # Standard-Template verwenden wenn keines angegeben
    if template_path is None:
        template_path = os.path.join(os.path.dirname(__file__), 'templates', 'solvengine.docx')

    if not os.path.exists(template_path):
        return f"Fehler: Template nicht gefunden: {template_path}"

    extra_args = ['--reference-doc', template_path]

    # Title page metadata - requires standalone mode
    title_page_enabled = False
    if doc_title:
        title_page_enabled = True
        extra_args.append('--standalone')
        extra_args.extend(['--metadata', f'title={doc_title}'])

        if doc_subtitle:
            extra_args.extend(['--metadata', f'subtitle={doc_subtitle}'])

        if doc_author:
            extra_args.extend(['--metadata', f'author={doc_author}'])

        if doc_date:
            extra_args.extend(['--metadata', f'date={doc_date}'])

    # Lua Filter für erweiterte Styling-Optionen (z.B. custom table styles)
    if lua_filter_path is None:
        # Standard-Filter für Tabellen-Styling verwenden
        default_filter = os.path.join(os.path.dirname(__file__), 'filters', 'table-style.lua')
        if os.path.exists(default_filter):
            lua_filter_path = default_filter

    if lua_filter_path and os.path.exists(lua_filter_path):
        extra_args.extend(['--lua-filter', lua_filter_path])

    # Table of Contents
    if enable_toc:
        extra_args.append('--toc')
        extra_args.extend(['--toc-depth', str(toc_depth)])
        if not doc_title:  # Only add standalone if not already added by title page
            extra_args.append('--standalone')  # Required for TOC generation in DOCX

    # Page Break Verarbeitung
    if process_page_breaks:
        extra_args.extend(['--from', 'html+raw_html'])

    # Input vorbereiten
    html_content = input_html
    if os.path.isfile(input_html):
        with open(input_html, 'r', encoding='utf-8') as f:
            html_content = f.read()

    # Page Breaks konvertieren
    if process_page_breaks:
        import re
        html_content = re.sub(
            r'<div[^>]*class=["\']page-?break["\'][^>]*>.*?</div>',
            '\n\n\\newpage\n\n',
            html_content,
            flags=re.IGNORECASE | re.DOTALL
        )
        html_content = re.sub(
            r'<hr[^>]*class=["\']page-?break["\'][^>]*>',
            '\n\n\\newpage\n\n',
            html_content,
            flags=re.IGNORECASE
        )
        html_content = re.sub(
            r'<div[^>]*style=["\'][^"\']*page-break-after\s*:\s*always[^"\']*["\'][^>]*>.*?</div>',
            '\n\n\\newpage\n\n',
            html_content,
            flags=re.IGNORECASE | re.DOTALL
        )

    # Konvertierung durchführen
    output = pypandoc.convert_text(
        html_content,
        'docx',
        format='html',
        outputfile=output_path,
        extra_args=extra_args
    )

    # Post-processing: Add cantSplit, page breaks, logo, etc.
    if prevent_table_split or (enable_toc and toc_page_break) or title_page_enabled or title_logo_path:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document(output_path)

        # Add logo to title page
        if title_page_enabled and title_logo_path and os.path.exists(title_logo_path):
            # Find the Title paragraph in body elements
            body_elements = list(doc.element.body)
            for i, elem in enumerate(body_elements):
                if elem.tag == qn('w:p'):
                    pPr = elem.find(qn('w:pPr'))
                    if pPr is not None:
                        pStyle = pPr.find(qn('w:pStyle'))
                        if pStyle is not None and pStyle.get(qn('w:val')) == 'Title':
                            # Create new paragraph element for logo
                            logo_p = OxmlElement('w:p')
                            logo_pPr = OxmlElement('w:pPr')
                            logo_jc = OxmlElement('w:jc')
                            logo_jc.set(qn('w:val'), 'center')
                            logo_pPr.append(logo_jc)
                            logo_p.append(logo_pPr)

                            # Insert after Title paragraph
                            doc.element.body.insert(i + 1, logo_p)

                            # Now use python-docx to add the picture
                            # Find the paragraph we just inserted
                            for para in doc.paragraphs:
                                if para._element == logo_p:
                                    run = para.add_run()
                                    run.add_picture(title_logo_path, width=Inches(title_logo_width_inches))
                                    break
                            break

        # Add page break after title page (before TOC or content)
        if title_page_enabled:
            # Find the Date paragraph element in body (last element of title page)
            body_elements = list(doc.element.body)
            for i, elem in enumerate(body_elements):
                # Check if this is a paragraph with Date style
                if elem.tag == qn('w:p'):
                    pPr = elem.find(qn('w:pPr'))
                    if pPr is not None:
                        pStyle = pPr.find(qn('w:pStyle'))
                        if pStyle is not None and pStyle.get(qn('w:val')) == 'Date':
                            # Next element (could be SDT or paragraph) should get a page break
                            if i + 1 < len(body_elements):
                                next_elem = body_elements[i + 1]

                                # If next element is SDT (like TOC), find first paragraph inside it
                                if next_elem.tag == qn('w:sdt'):
                                    # Find first paragraph in SDT content
                                    first_para_in_sdt = next_elem.find(f'.//{qn("w:p")}')
                                    if first_para_in_sdt is not None:
                                        target_para = first_para_in_sdt
                                        pPr_target = target_para.find(qn('w:pPr'))
                                        if pPr_target is None:
                                            pPr_target = OxmlElement('w:pPr')
                                            target_para.insert(0, pPr_target)
                                        pageBreakBefore = OxmlElement('w:pageBreakBefore')
                                        pPr_target.append(pageBreakBefore)
                                # If next element is a paragraph
                                elif next_elem.tag == qn('w:p'):
                                    pPr_target = next_elem.find(qn('w:pPr'))
                                    if pPr_target is None:
                                        pPr_target = OxmlElement('w:pPr')
                                        next_elem.insert(0, pPr_target)
                                    pageBreakBefore = OxmlElement('w:pageBreakBefore')
                                    pPr_target.append(pageBreakBefore)
                            break

        # Add cantSplit to tables
        if prevent_table_split:
            for table in doc.tables:
                for row in table.rows:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    cantSplit = OxmlElement('w:cantSplit')
                    trPr.append(cantSplit)

        # Add page break after TOC
        if enable_toc and toc_page_break:
            # Find the TOC SDT (Structured Document Tag) block
            sdt_blocks = doc.element.body.xpath('.//w:sdt')

            for sdt in sdt_blocks:
                # Check if this SDT contains a TOC
                doc_part_gallery = sdt.xpath('.//w:docPartGallery', namespaces=sdt.nsmap)
                if doc_part_gallery and doc_part_gallery[0].get(qn('w:val')) == 'Table of Contents':
                    # Find the next paragraph after the SDT block
                    next_elem = sdt.getnext()

                    # Skip bookmarks and find the actual paragraph
                    while next_elem is not None and next_elem.tag != qn('w:p'):
                        next_elem = next_elem.getnext()

                    if next_elem is not None and next_elem.tag == qn('w:p'):
                        # Insert page break at the beginning of this paragraph
                        pPr = next_elem.find(qn('w:pPr'))
                        if pPr is None:
                            pPr = OxmlElement('w:pPr')
                            next_elem.insert(0, pPr)

                        # Add page break before
                        pageBreakBefore = OxmlElement('w:pageBreakBefore')
                        pPr.append(pageBreakBefore)
                        break

        doc.save(output_path)

    logging.info(f"Converted HTML to DOCX with template: {output_path}")
    return f"Erfolgreich konvertiert mit Template nach: {output_path}"


@mcp.tool()
def create_docx_with_lua_filter(
    input_html: str,
    output_path: str,
    lua_filter_path: str,
    header_text: str = None,
    footer_text: str = None
) -> str:
    """
    Konvertiert HTML zu Word (DOCX) mit Lua-Filter für erweiterte Header/Footer-Kontrolle.

    Args:
        input_html: HTML-String oder Pfad zur HTML-Datei
        output_path: Pfad für die Ausgabe-DOCX-Datei
        lua_filter_path: Pfad zum Lua-Filter für Header/Footer
        header_text: Optionaler Header-Text (wird an Filter übergeben)
        footer_text: Optionaler Footer-Text (wird an Filter übergeben)

    Returns:
        Bestätigung mit Dateipfad
    """
    if not os.path.exists(lua_filter_path):
        return f"Fehler: Lua-Filter nicht gefunden: {lua_filter_path}"

    extra_args = ['--lua-filter', lua_filter_path]

    # Header/Footer als Variablen übergeben
    if header_text:
        extra_args.extend(['--variable', f'header={header_text}'])

    if footer_text:
        extra_args.extend(['--variable', f'footer={footer_text}'])

    # Prüfen ob input_html ein Dateipfad ist
    if os.path.isfile(input_html):
        output = pypandoc.convert_file(
            input_html,
            'docx',
            outputfile=output_path,
            extra_args=extra_args
        )
    else:
        # Als HTML-String behandeln
        output = pypandoc.convert_text(
            input_html,
            'docx',
            format='html',
            outputfile=output_path,
            extra_args=extra_args
        )

    logging.info(f"Converted HTML to DOCX with Lua filter: {output_path}")
    return f"Erfolgreich konvertiert mit Lua-Filter nach: {output_path}"


@mcp.tool()
def convert_md_to_docx(
    input_markdown: str,
    output_path: str,
    header_text: str = None,
    footer_text: str = None,
    reference_doc: str = None,
    enable_toc: bool = True,
    toc_depth: int = 3,
    lua_filter_path: str = None,
    prevent_table_split: bool = True,
    toc_page_break: bool = True,
    doc_title: str = None,
    doc_subtitle: str = None,
    doc_author: str = None,
    doc_date: str = None,
    title_logo_path: str = None,
    title_logo_width_inches: float = 2.0
) -> str:
    """
    Konvertiert Markdown zu Word (DOCX) mit optionalen Header/Footer, TOC und Styling.

    Args:
        input_markdown: Markdown-String oder Pfad zur Markdown-Datei
        output_path: Pfad für die Ausgabe-DOCX-Datei
        header_text: Optionaler Header-Text
        footer_text: Optionaler Footer-Text
        reference_doc: Pfad zu einer Referenz-DOCX-Datei für Styling (Standard: templates/solvengine.docx)
        enable_toc: Aktiviert automatisches Inhaltsverzeichnis
        toc_depth: Tiefe des Inhaltsverzeichnisses (1-6, Standard: 3)
        lua_filter_path: Pfad zu einem Lua-Filter für erweiterte Styling-Optionen
        prevent_table_split: Verhindert dass Tabellen über Seitengrenzen getrennt werden
        toc_page_break: Fügt einen Seitenumbruch nach dem Inhaltsverzeichnis ein
        doc_title: Dokumenttitel für Titelseite
        doc_subtitle: Untertitel für Titelseite
        doc_author: Autor(en) für Titelseite
        doc_date: Datum für Titelseite (wenn None, wird aktuelles Datum verwendet)
        title_logo_path: Pfad zum Logo-Bild für Titelseite
        title_logo_width_inches: Breite des Logos in Zoll (Standard: 2.0)

    Returns:
        Bestätigung mit Dateipfad
    """
    extra_args = []

    # Title page metadata - requires standalone mode
    title_page_enabled = False
    if doc_title:
        title_page_enabled = True
        extra_args.append('--standalone')
        extra_args.extend(['--metadata', f'title={doc_title}'])

        if doc_subtitle:
            extra_args.extend(['--metadata', f'subtitle={doc_subtitle}'])

        if doc_author:
            extra_args.extend(['--metadata', f'author={doc_author}'])

        if doc_date:
            extra_args.extend(['--metadata', f'date={doc_date}'])

    # Header und Footer als Metadaten hinzufügen
    if header_text:
        extra_args.extend(['--metadata', f'header-includes={header_text}'])

    if footer_text:
        extra_args.extend(['--metadata', f'footer-includes={footer_text}'])

    # Referenz-Dokument für Styling
    if reference_doc is None:
        # Standard-Template verwenden
        default_template = os.path.join(os.path.dirname(__file__), 'templates', 'solvengine.docx')
        if os.path.exists(default_template):
            reference_doc = default_template

    if reference_doc and os.path.exists(reference_doc):
        extra_args.extend(['--reference-doc', reference_doc])

    # Lua Filter
    if lua_filter_path is None:
        # Standard-Filter für Tabellen-Styling verwenden
        default_filter = os.path.join(os.path.dirname(__file__), 'filters', 'table-style.lua')
        if os.path.exists(default_filter):
            lua_filter_path = default_filter

    if lua_filter_path and os.path.exists(lua_filter_path):
        extra_args.extend(['--lua-filter', lua_filter_path])

    # Table of Contents
    if enable_toc:
        extra_args.append('--toc')
        extra_args.extend(['--toc-depth', str(toc_depth)])
        if not doc_title:  # Only add standalone if not already added by title page
            extra_args.append('--standalone')  # Required for TOC generation in DOCX

    # Input vorbereiten
    md_content = input_markdown
    if os.path.isfile(input_markdown):
        with open(input_markdown, 'r', encoding='utf-8') as f:
            md_content = f.read()

    # Konvertierung durchführen
    output = pypandoc.convert_text(
        md_content,
        'docx',
        format='markdown',
        outputfile=output_path,
        extra_args=extra_args
    )

    # Post-processing: Add cantSplit, page breaks, logo, etc.
    if prevent_table_split or (enable_toc and toc_page_break) or title_page_enabled or title_logo_path:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document(output_path)

        # Add logo to title page
        if title_page_enabled and title_logo_path and os.path.exists(title_logo_path):
            # Find the Title paragraph in body elements
            body_elements = list(doc.element.body)
            for i, elem in enumerate(body_elements):
                if elem.tag == qn('w:p'):
                    pPr = elem.find(qn('w:pPr'))
                    if pPr is not None:
                        pStyle = pPr.find(qn('w:pStyle'))
                        if pStyle is not None and pStyle.get(qn('w:val')) == 'Title':
                            # Create new paragraph element for logo
                            logo_p = OxmlElement('w:p')
                            logo_pPr = OxmlElement('w:pPr')
                            logo_jc = OxmlElement('w:jc')
                            logo_jc.set(qn('w:val'), 'center')
                            logo_pPr.append(logo_jc)
                            logo_p.append(logo_pPr)

                            # Insert after Title paragraph
                            doc.element.body.insert(i + 1, logo_p)

                            # Now use python-docx to add the picture
                            # Find the paragraph we just inserted
                            for para in doc.paragraphs:
                                if para._element == logo_p:
                                    run = para.add_run()
                                    run.add_picture(title_logo_path, width=Inches(title_logo_width_inches))
                                    break
                            break

        # Add page break after title page (before TOC or content)
        if title_page_enabled:
            # Find the Date paragraph element in body (last element of title page)
            body_elements = list(doc.element.body)
            for i, elem in enumerate(body_elements):
                # Check if this is a paragraph with Date style
                if elem.tag == qn('w:p'):
                    pPr = elem.find(qn('w:pPr'))
                    if pPr is not None:
                        pStyle = pPr.find(qn('w:pStyle'))
                        if pStyle is not None and pStyle.get(qn('w:val')) == 'Date':
                            # Next element (could be SDT or paragraph) should get a page break
                            if i + 1 < len(body_elements):
                                next_elem = body_elements[i + 1]

                                # If next element is SDT (like TOC), find first paragraph inside it
                                if next_elem.tag == qn('w:sdt'):
                                    # Find first paragraph in SDT content
                                    first_para_in_sdt = next_elem.find(f'.//{qn("w:p")}')
                                    if first_para_in_sdt is not None:
                                        target_para = first_para_in_sdt
                                        pPr_target = target_para.find(qn('w:pPr'))
                                        if pPr_target is None:
                                            pPr_target = OxmlElement('w:pPr')
                                            target_para.insert(0, pPr_target)
                                        pageBreakBefore = OxmlElement('w:pageBreakBefore')
                                        pPr_target.append(pageBreakBefore)
                                # If next element is a paragraph
                                elif next_elem.tag == qn('w:p'):
                                    pPr_target = next_elem.find(qn('w:pPr'))
                                    if pPr_target is None:
                                        pPr_target = OxmlElement('w:pPr')
                                        next_elem.insert(0, pPr_target)
                                    pageBreakBefore = OxmlElement('w:pageBreakBefore')
                                    pPr_target.append(pageBreakBefore)
                            break

        # Add cantSplit to tables
        if prevent_table_split:
            for table in doc.tables:
                for row in table.rows:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    cantSplit = OxmlElement('w:cantSplit')
                    trPr.append(cantSplit)

        # Add page break after TOC
        if enable_toc and toc_page_break:
            # Find the TOC SDT (Structured Document Tag) block
            sdt_blocks = doc.element.body.xpath('.//w:sdt')

            for sdt in sdt_blocks:
                # Check if this SDT contains a TOC
                doc_part_gallery = sdt.xpath('.//w:docPartGallery', namespaces=sdt.nsmap)
                if doc_part_gallery and doc_part_gallery[0].get(qn('w:val')) == 'Table of Contents':
                    # Find the next paragraph after the SDT block
                    next_elem = sdt.getnext()

                    # Skip bookmarks and find the actual paragraph
                    while next_elem is not None and next_elem.tag != qn('w:p'):
                        next_elem = next_elem.getnext()

                    if next_elem is not None and next_elem.tag == qn('w:p'):
                        # Insert page break at the beginning of this paragraph
                        pPr = next_elem.find(qn('w:pPr'))
                        if pPr is None:
                            pPr = OxmlElement('w:pPr')
                            next_elem.insert(0, pPr)

                        # Add page break before
                        pageBreakBefore = OxmlElement('w:pageBreakBefore')
                        pPr.append(pageBreakBefore)
                        break

        doc.save(output_path)

    logging.info(f"Converted Markdown to DOCX: {output_path}")
    return f"Erfolgreich konvertiert nach: {output_path}"


@mcp.tool()
def check_pandoc_version() -> str:
    """
    Prüft ob Pandoc installiert ist und gibt die Version zurück.

    Returns:
        Pandoc-Version
    """
    version = pypandoc.get_pandoc_version()
    logging.info(f"Pandoc version: {version}")
    return f"Pandoc Version: {version}"


@mcp.tool()
def list_pandoc_formats() -> str:
    """
    Listet alle verfügbaren Pandoc-Eingabe- und Ausgabeformate.

    Returns:
        Dictionary mit Input/Output-Formaten
    """
    formats = pypandoc.get_pandoc_formats()
    logging.info(f"Pandoc formats: {formats}")
    return str(formats)


# -----------------------------------------------------------------------------
# Start MCP Server
# -----------------------------------------------------------------------------
if __name__ == "__main__":
    mcp.run()
